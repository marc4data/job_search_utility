"""
build_docs.py — resume + cover-letter builder and ATS-critique engine.

This engine contains NO personal data. Everything person-specific is read
from profile/profile.py (the PROFILE dict). Base resumes live in
profile/bases.py (the BASES dict) and are passed in by the batch script.

Generated for the `job-search-tailor` plugin. Two important invariants:
  • critique_and_refine() is a CONTENT-QA tool only and NEVER writes to the
    tracker. The True ATS score (written to the tracker) is produced solely
    by the separate ats_score script.
  • Never claim a skill/tool that is not in profile/verified_skills.md.
"""
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    import sys
    sys.exit(
        f"\n[job-search-tailor] Missing dependency: {e.name}\n"
        "The engine needs python-docx (and openpyxl for scoring).\n"
        "Install them into the SAME Python you run scripts with:\n\n"
        "    python3 -m pip install -r requirements.txt\n\n"
        "Use 'python3 -m pip' (not bare 'pip') so the installer and the\n"
        "interpreter that runs this engine are the same one.\n"
    )
from datetime import date
import os

import importlib.util

def _load_profile():
    """Load the user's PROFILE dict from profile/profile.py.

    Searched, in order: env JS_PROFILE, ./profile/profile.py, ../profile/profile.py,
    and next to this engine file. Run the `setup-profile` skill to generate it.
    """
    here = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.environ.get("JS_PROFILE", ""),
        os.path.join(os.getcwd(), "profile", "profile.py"),
        os.path.join(here, "..", "profile", "profile.py"),
        os.path.join(here, "profile.py"),
    ]
    for c in candidates:
        if c and os.path.exists(c):
            spec = importlib.util.spec_from_file_location("js_profile", c)
            mod = importlib.util.module_from_spec(spec); spec.loader.exec_module(mod)
            return mod.PROFILE
    raise FileNotFoundError(
        "profile/profile.py not found. Run the setup-profile skill first.")

PROFILE = _load_profile()

OUT_DIR = os.environ.get("JS_OUT_DIR", os.getcwd())
os.makedirs(OUT_DIR, exist_ok=True)

NAVY  = RGBColor(0x1F, 0x38, 0x64)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x40, 0x40, 0x40)

RIGHT_TAB = 10080  # twips = full 7" text width (8.5" page - 0.75" margins each side)

# ── pPr element ordering (OOXML schema) ──────────────────────────────────────
# Correct child-element order within w:pPr (subset we use):
PPR_ORDER = [
    "pStyle","keepNext","keepLines","pageBreakBefore","framePr","widowControl",
    "numPr","suppressLineNumbers","pBdr","shd","tabs","suppressAutoHyphens",
    "kinsoku","wordWrap","overflowPunct","topLinePunct","autoSpaceDE",
    "autoSpaceDN","bidi","adjustRightInd","snapToGrid","spacing","ind",
    "contextualSpacing","mirrorIndents","suppressOverlap","jc","textDirection",
    "textAlignment","outlineLvl","rPr","sectPr","pPrChange",
]

def _ppr_insert(pPr, new_el):
    """Insert an element into pPr in schema order."""
    tag = new_el.tag.split("}")[-1] if "}" in new_el.tag else new_el.tag
    try:
        new_pos = PPR_ORDER.index(tag)
    except ValueError:
        pPr.append(new_el)
        return
    # Find first existing child whose position is AFTER new_pos and insert before it
    for child in list(pPr):
        child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        try:
            child_pos = PPR_ORDER.index(child_tag)
        except ValueError:
            child_pos = 999
        if child_pos > new_pos:
            child.addprevious(new_el)
            return
    pPr.append(new_el)


def set_page_us_letter(doc):
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)

def add_bottom_border(para, color="2E75B6", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    _ppr_insert(pPr, pBdr)

def add_right_tab_stop(para, pos=RIGHT_TAB):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)
    _ppr_insert(pPr, tabs)

def set_space_before_after(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    # Remove existing spacing element first to avoid duplicates
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    _ppr_insert(pPr, spacing)

def set_indent(para, left=0, hanging=0):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:ind")):
        pPr.remove(existing)
    ind = OxmlElement("w:ind")
    if left:    ind.set(qn("w:left"),    str(left))
    if hanging: ind.set(qn("w:hanging"), str(hanging))
    _ppr_insert(pPr, ind)


# ─────────────────────────────────────────────────────────────────────────────
# Content helpers
# ─────────────────────────────────────────────────────────────────────────────

def add_name_header(doc, name, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space_before_after(p, 0, 20)
    r = p.add_run(name)
    r.font.name = "Arial"; r.font.size = Pt(20)
    r.font.bold = True; r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space_before_after(p2, 0, 10)
    r2 = p2.add_run(title)
    r2.font.name = "Arial"; r2.font.size = Pt(12)
    r2.font.color.rgb = BLUE

def _add_hyperlink(para, text, url, size=9, color="2E75B6"):
    """Append a clickable hyperlink run to an existing paragraph."""
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, attrs in [
        ("w:rFonts", {qn("w:ascii"): "Arial", qn("w:hAnsi"): "Arial"}),
        ("w:sz",     {qn("w:val"): str(size * 2)}),
        ("w:szCs",   {qn("w:val"): str(size * 2)}),
        ("w:color",  {qn("w:val"): color}),
        ("w:u",      {qn("w:val"): "single"}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(k, v)
        rPr.append(el)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def _add_plain_run(para, text, size=9, color=None):
    """Append a plain (non-hyperlink) run to an existing paragraph."""
    r = para.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = color or GRAY


def add_contact_line(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_space_before_after(p, 0, 6)
    add_bottom_border(p, "2E75B6", 8)

    _add_plain_run(p, PROFILE["contact"])
    for label, url in PROFILE.get("links", []):
        _add_plain_run(p, "  \u2022  ")
        _add_hyperlink(p, label, url)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    set_space_before_after(p, 100, 30)
    add_bottom_border(p, "2E75B6", 8)
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.bold = True; r.font.color.rgb = NAVY
    return p

def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    set_space_before_after(p, 0, 20)
    set_indent(p, left=220, hanging=220)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    return p

def add_experience_header(doc, role, company_loc, date_range):
    # Line 1 — Role title (bold, navy)
    p1 = doc.add_paragraph()
    set_space_before_after(p1, 90, 0)
    r1 = p1.add_run(role)
    r1.font.name = "Arial"; r1.font.size = Pt(10.5)
    r1.font.bold = True; r1.font.color.rgb = NAVY

    # Line 2 — Company | Location [tab right] Date
    p2 = doc.add_paragraph()
    set_space_before_after(p2, 0, 0)
    add_right_tab_stop(p2, RIGHT_TAB)

    r2 = p2.add_run(company_loc)
    r2.font.name = "Arial"; r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY

    p2.add_run("\t")

    r3 = p2.add_run(date_range)
    r3.font.name = "Arial"; r3.font.size = Pt(10)
    r3.font.italic = True; r3.font.color.rgb = BLUE
    return p1

def add_company_subhead(doc, text):
    p = doc.add_paragraph()
    set_space_before_after(p, 0, 30)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(9.5)
    r.font.italic = True; r.font.color.rgb = GRAY
    return p

def add_plain(doc, text, size=9.5, bold=False, color=None, before=0, after=30):
    p = doc.add_paragraph()
    set_space_before_after(p, before, after)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or BLACK
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Resume data
# ─────────────────────────────────────────────────────────────────────────────

# ---------------------------------------------------------------------------
# Base-resume DATA lives in profile/bases.py (BASES dict), generated by the
# setup-profile skill and passed in by the batch script. None is stored here.
# ---------------------------------------------------------------------------

# =============================================================================

def build_resume(filename, target_title, summary_lines, expertise_items,
                 career_highlights=None, experience=None,
                 projects=None, tech_expertise=None, certs_before_edu=False):
    doc = Document()
    set_page_us_letter(doc)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    # Use overrides if provided, else fall back to original globals
    _highlights   = career_highlights if career_highlights is not None else []
    _experience   = experience        if experience        is not None else []
    _tech         = tech_expertise    if tech_expertise    is not None else []

    add_name_header(doc, PROFILE["name"].upper(), target_title)
    add_contact_line(doc)

    add_section_header(doc, "Professional Summary")
    for line in summary_lines:
        add_plain(doc, line, size=9.5, after=50)

    add_section_header(doc, "Areas of Expertise")
    p = doc.add_paragraph()
    set_space_before_after(p, 0, 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("  |  ".join(expertise_items))
    r.font.name = "Arial"; r.font.size = Pt(9)
    r.font.color.rgb = NAVY

    add_section_header(doc, "Career Highlights")
    for h in _highlights:
        add_bullet(doc, h)

    add_section_header(doc, "Professional Experience")
    for exp in _experience:
        add_experience_header(doc, exp["role"], exp["company_loc"], exp["dates"])
        add_company_subhead(doc, exp["desc"])
        for b in exp["bullets"]:
            add_bullet(doc, b)

    # Optional Projects section (IC base only)
    if projects:
        add_section_header(doc, "Selected Projects")
        for proj in projects:
            p_hdr = doc.add_paragraph()
            set_space_before_after(p_hdr, 80, 0)
            r_name = p_hdr.add_run(proj["name"] + "  ")
            r_name.font.name = "Arial"; r_name.font.size = Pt(10.5)
            r_name.font.bold = True; r_name.font.color.rgb = NAVY
            r_tech = p_hdr.add_run(proj["tech"])
            r_tech.font.name = "Arial"; r_tech.font.size = Pt(9)
            r_tech.font.italic = True; r_tech.font.color.rgb = GRAY
            for b in proj["bullets"]:
                add_bullet(doc, b)

    if PROFILE.get("additional_experience"):
        add_section_header(doc, "Additional Experience")
        for a in PROFILE["additional_experience"]:
            add_plain(doc, a, size=9.5, after=30)

    # IC: put certifications BEFORE education so technical creds land higher
    _edu   = PROFILE.get("education", [])
    _certs = PROFILE.get("certifications", [])

    if certs_before_edu:
        add_section_header(doc, "Certifications & Training")
        for c in _certs:
            add_plain(doc, c, size=9.5, after=30)
        add_section_header(doc, "Education")
        for e in _edu:
            add_plain(doc, e, size=9.5, after=30)
    else:
        add_section_header(doc, "Education")
        for e in _edu:
            add_plain(doc, e, size=9.5, after=30)
        add_section_header(doc, "Certifications & Training")
        for c in _certs:
            add_plain(doc, c, size=9.5, after=30)

    add_section_header(doc, "Technical Expertise")
    for label, items in _tech:
        p = doc.add_paragraph()
        set_space_before_after(p, 0, 30)
        rb = p.add_run(label + ":  ")
        rb.font.name = "Arial"; rb.font.size = Pt(9.5)
        rb.font.bold = True; rb.font.color.rgb = NAVY
        rr = p.add_run(items)
        rr.font.name = "Arial"; rr.font.size = Pt(9.5)
        rr.font.color.rgb = BLACK

    doc.save(filename)
    print(f"  \u2713 {os.path.basename(filename)}")


# ─────────────────────────────────────────────────────────────────────────────
# Cover letter builder
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_letter(filename, company, role, location, body_paragraphs):
    doc = Document()
    set_page_us_letter(doc)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    today_str = date.today().strftime("%B %d, %Y")

    add_plain(doc, today_str, size=10.5, after=60)
    for line in ["Hiring Manager", company, location]:
        add_plain(doc, line, size=10.5, after=20)
    doc.add_paragraph()  # blank line
    add_plain(doc, "Dear Hiring Manager,", size=10.5, after=80)

    p = doc.add_paragraph()
    set_space_before_after(p, 0, 120)
    r = p.add_run("Re: " + role)
    r.font.name = "Arial"; r.font.size = Pt(10.5)
    r.font.bold = True; r.font.color.rgb = NAVY

    for para_text in body_paragraphs:
        p = doc.add_paragraph()
        set_space_before_after(p, 0, 140)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(para_text)
        r.font.name = "Arial"; r.font.size = Pt(10.5)
        r.font.color.rgb = BLACK

    doc.add_paragraph()
    add_plain(doc, "Sincerely,", size=10.5, after=260)

    p = doc.add_paragraph()
    set_space_before_after(p, 0, 20)
    r = p.add_run(PROFILE["name"])
    r.font.name = "Arial"; r.font.size = Pt(10.5)
    r.font.bold = True; r.font.color.rgb = NAVY

    add_plain(doc, PROFILE.get("cl_contact", PROFILE["contact"]),
              size=10, color=GRAY, after=0)

    doc.save(filename)
    print(f"  \u2713 {os.path.basename(filename)}")


# =============================================================================
# ATS SCORING  —  score_and_log() is the only function batch scripts need
# =============================================================================

TRACKER_PATH = os.environ.get("JS_TRACKER_PATH", "")  # legacy helper only; ATS script writes the tracker

def ats_score_resume(resume_path, jd_keywords):
    """
    Score a resume .docx against a keyword list.

    jd_keywords accepts two formats:
      • Flat list of strings          → ["dbt", "Snowflake", "Python", ...]
      • Weighted tuples               → [("dbt", 3), ("Python", 2), ("agile", 1), ...]

    Weights allow you to mark required skills heavier than nice-to-haves:
      3 = must-have  |  2 = strong differentiator  |  1 = contextual / soft skill

    Returns: integer 0-100
    """
    from docx import Document as _D
    doc  = _D(resume_path)
    text = " ".join(p.text for p in doc.paragraphs).lower()
    # include any table cells (expertise strip is a paragraph, but be safe)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text += " " + cell.text.lower()

    if not jd_keywords:
        return 0

    # Normalise to (keyword, weight) tuples
    if isinstance(jd_keywords[0], (list, tuple)):
        weighted = [(str(kw).lower(), int(w)) for kw, w in jd_keywords]
    else:
        weighted = [(str(kw).lower(), 1) for kw in jd_keywords]

    max_pts = sum(w for _, w in weighted)
    hit_pts = sum(w for kw, w in weighted if kw in text)
    return round((hit_pts / max_pts) * 100) if max_pts else 0


def _write_tracker_score(company, role, score):
    """
    Write score to column K (Resume Score) in the tracker.
    Matches on first word of company + first 3 words of role title.
    Returns True if the row was found and updated.
    """
    from openpyxl import load_workbook as _lw
    wb  = _lw(TRACKER_PATH)
    ws  = wb["📋 Applications"]
    hdrs = [c.value for c in ws[3]]

    co_col    = hdrs.index("Company")      + 1
    role_col  = hdrs.index("Role / Title") + 1
    score_col = hdrs.index("Resume Score") + 1

    # Normalise search terms
    co_stem   = company.lower().replace(",", "").split()[0][:6]   # first 6 chars of first word
    role_words = [w for w in role.lower().split()[:4]]             # first 4 role words

    for row in ws.iter_rows(min_row=4):
        cell_co   = str(row[co_col   - 1].value or "").lower()
        cell_role = str(row[role_col - 1].value or "").lower()
        co_hit    = co_stem in cell_co
        role_hit  = sum(1 for w in role_words if w in cell_role) >= 2
        if co_hit and role_hit:
            row[score_col - 1].value = score
            wb.save(TRACKER_PATH)
            return True
    return False


def score_and_log(resume_path, company, role, jd_keywords):
    """
    One-liner for batch scripts:
        score_and_log(resume_path, "Acme Corp", "Senior Analytics Engineer", jd_keywords)

    Scores the resume, prints result, writes to tracker.
    Returns the integer score.
    """
    score = ats_score_resume(resume_path, jd_keywords)
    found = _write_tracker_score(company, role, score)
    tag   = "✓ logged" if found else "⚠ tracker row not matched"
    print(f"  ATS Score: {score}/100  [{tag}]")
    return score


# =============================================================================
# TWO-PASS REVIEW SYSTEM
# =============================================================================
#
# critique_resume()     — honest second-perspective scoring with word-boundary
#                         matching and domain-specific hard terms.  Produces a
#                         gap report: what the JD requires that the draft lacks.
#
# apply_gap_patches()   — applies targeted text patches to summary + expertise
#                         for gaps that can be honestly addressed.  Does NOT
#                         fabricate credentials; patches should reflect real
#                         experience described in natural domain language.
#
# critique_and_refine() — the single call batch scripts use.  Runs critique,
#                         patches where possible, rebuilds if needed, re-scores,
#                         and logs the final honest score.
#
# Batch script pattern:
#
#   role_1["jd_hard"] = [("pharma",3), ("clinical",2), ("veeva",2), ...]
#   role_1["gap_patches"] = {
#       "pharma":   {"target": "p2",        "text": " in pharma and life sciences"},
#       "clinical": {"target": "expertise", "text": "Life Sciences / Clinical Data"},
#       "veeva":    None,   # genuine gap — flag, don't patch
#   }
#   critique_and_refine(
#       resume_path, company, role, target_title,
#       summary_lines, expertise_items,
#       role_dict,               # must have "jd_hard" and "gap_patches" keys
#       career_highlights=..., experience=..., tech_expertise=...,
#   )
# =============================================================================

import re as _re

def critique_resume(resume_path, jd_hard):
    """
    Honest second-perspective scoring.

    jd_hard — list of (keyword, weight) tuples that include domain-specific
               terms the JD actually requires, not just generic base terms.
               Uses word-boundary matching so 'sql' doesn't fire on 'mysql'.

    Returns dict:
        score     — int 0-100
        hits      — [(kw, weight), ...]
        misses    — [(kw, weight), ...]  sorted heaviest first
        hit_pts   — int
        max_pts   — int
        critical  — [kw, ...]  misses with weight >= 3
    """
    from docx import Document as _D
    doc  = _D(resume_path)
    text = " ".join(p.text for p in doc.paragraphs).lower()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text += " " + cell.text.lower()

    if not jd_hard:
        return {"score": 0, "hits": [], "misses": [], "hit_pts": 0, "max_pts": 0, "critical": []}

    normalized = [(str(kw).lower(), int(w)) for kw, w in jd_hard]
    max_pts  = sum(w for _, w in normalized)

    hits   = [(kw, w) for kw, w in normalized
              if _re.search(r'\b' + _re.escape(kw) + r'\b', text)]
    misses = [(kw, w) for kw, w in normalized
              if not _re.search(r'\b' + _re.escape(kw) + r'\b', text)]
    misses.sort(key=lambda x: -x[1])

    hit_pts  = sum(w for _, w in hits)
    score    = round((hit_pts / max_pts) * 100) if max_pts else 0
    critical = [kw for kw, w in misses if w >= 3]

    return {
        "score":    score,
        "hits":     hits,
        "misses":   misses,
        "hit_pts":  hit_pts,
        "max_pts":  max_pts,
        "critical": critical,
    }


def apply_gap_patches(summary_lines, expertise_items, misses, gap_patches):
    """
    Apply targeted text additions for addressable gaps.

    gap_patches dict format:
        {
          "pharma": {"target": "p2",        "text": " in pharma and life sciences"},
          "airflow": {"target": "expertise", "text": "Airflow"},
          "veeva":   None,   # genuine gap — keep as None to flag but skip patching
        }

    targets:
        "p1"        — appends text to summary paragraph 1
        "p2"        — appends text to summary paragraph 2  (most common for domain/tech)
        "p3"        — appends text to summary paragraph 3
        "expertise" — adds a new item to the expertise list (if not already present)

    Returns: (patched_summary_lines, patched_expertise_items, applied, skipped)
        applied — list of keywords that were patched
        skipped — list of keywords with no patch defined (honest gaps)
    """
    summary = list(summary_lines)
    expertise = list(expertise_items)
    applied = []
    skipped = []

    for kw, _w in misses:
        patch = gap_patches.get(kw)
        if patch is None:
            skipped.append(kw)
            continue
        target = patch.get("target", "p2")
        text   = patch.get("text", "")
        if not text:
            skipped.append(kw)
            continue

        if target == "expertise":
            if text not in expertise:
                expertise.append(text)
        elif target in ("p1", "p2", "p3"):
            idx = {"p1": 0, "p2": 1, "p3": 2}.get(target, 1)
            if idx < len(summary):
                # Only append if the patch text isn't already present
                if text.lower().strip() not in summary[idx].lower():
                    summary[idx] = summary[idx].rstrip(". ") + text
        applied.append(kw)

    return summary, expertise, applied, skipped


def apply_tight_patches(tech_expertise, remaining_misses, tight_patches):
    """
    Pass 3: Precision keyword insertion into existing tech_expertise rows.

    tight_patches dict format:
        {
          "ci/cd":        {"row": "Data Engineering", "text": "CI/CD"},
          "self-service": {"row": "Analytic Tools",   "text": "Self-Service Analytics"},
          "databricks":   None,  # genuine gap — flag but don't insert
        }

    Finds the named row in tech_expertise (case-insensitive label match) and
    appends the text to that row's comma-delimited item string.  Only appends
    if the term isn't already present (case-insensitive check).

    Returns: (patched_tech_expertise, applied, skipped)
        applied — keywords whose term was inserted
        skipped — genuine gaps (None patch) or unmatched rows
    """
    te      = [list(row) for row in tech_expertise]
    applied = []
    skipped = []

    for kw, _w in remaining_misses:
        patch = tight_patches.get(kw)
        if patch is None:
            skipped.append(kw)
            continue
        row_label = patch.get("row", "")
        text      = patch.get("text", "")
        if not row_label or not text:
            skipped.append(kw)
            continue

        patched = False
        for row in te:
            if row[0].lower() == row_label.lower():
                current = row[1]
                if text.lower() not in current.lower():
                    row[1] = current.rstrip(", ") + ", " + text
                applied.append(kw)
                patched = True
                break
        if not patched:
            skipped.append(kw)

    return [tuple(row) for row in te], applied, skipped


def critique_and_refine(
    resume_path, company, role_title, target_title,
    summary_lines, expertise_items, role_dict,
    career_highlights=None, experience=None,
    tech_expertise=None, certs_before_edu=False,
    refine_threshold=80,
):
    """
    Three-pass review + refinement for batch scripts.

    Pass 1: critique the just-built draft against role_dict["jd_hard"]
    Pass 2: if score < refine_threshold OR critical gaps exist:
              apply role_dict["gap_patches"] (summary/expertise additions), rebuild
    Pass 3: apply role_dict["tight_patches"] (precision tech_expertise row
              augmentation) against remaining misses — always runs if provided
    Final:  log the honest score from the last critique performed

    role_dict must have:
        "jd_hard"       — list of (kw, weight) hard domain terms
        "gap_patches"   — dict as described in apply_gap_patches()
        "tight_patches" — dict as described in apply_tight_patches() [optional]
        "company"       — str
        "role"          — str

    Returns final honest score (int).
    """
    jd_hard       = role_dict.get("jd_hard", [])
    gap_patches   = role_dict.get("gap_patches", {})
    tight_patches = role_dict.get("tight_patches", {})

    # Working copies so originals stay clean
    cur_summary   = list(summary_lines)
    cur_expertise = list(expertise_items)
    cur_tech      = list(tech_expertise) if tech_expertise is not None else None

    # --- Pass 1: critique the draft ---
    c1 = critique_resume(resume_path, jd_hard)
    print(f"  Reviewer Pass 1: {c1['score']}/100", end="")
    if c1["critical"]:
        print(f"  ⚠ critical gaps: {c1['critical']}", end="")
    if c1["misses"]:
        miss_str = ", ".join(f"'{kw}'[{w}]" for kw, w in c1["misses"])
        print(f"\n    misses: {miss_str}", end="")
    print()

    current_critique = c1

    # --- Pass 2: strategic gap patches (summary / expertise additions) ---
    needs_refine = (c1["score"] < refine_threshold) or bool(c1["critical"])

    if needs_refine and gap_patches:
        patched_summary, patched_expertise, applied, skipped = apply_gap_patches(
            cur_summary, cur_expertise, c1["misses"], gap_patches
        )
        if applied:
            cur_summary   = patched_summary
            cur_expertise = patched_expertise
            build_resume(
                resume_path, target_title,
                cur_summary, cur_expertise,
                career_highlights=career_highlights,
                experience=experience,
                tech_expertise=cur_tech,
                certs_before_edu=certs_before_edu,
            )
            c2 = critique_resume(resume_path, jd_hard)
            print(f"  Reviewer Pass 2: {c2['score']}/100  (patched: {applied})", end="")
            if skipped:
                print(f"  | honest gaps remain: {skipped}", end="")
            print()
            current_critique = c2

    # --- Pass 3: precision tight patches (tech_expertise row augmentation) ---
    if tight_patches and current_critique["misses"] and cur_tech is not None:
        patched_tech, applied3, skipped3 = apply_tight_patches(
            cur_tech, current_critique["misses"], tight_patches
        )
        if applied3:
            cur_tech = patched_tech
            build_resume(
                resume_path, target_title,
                cur_summary, cur_expertise,
                career_highlights=career_highlights,
                experience=experience,
                tech_expertise=cur_tech,
                certs_before_edu=certs_before_edu,
            )
            c3 = critique_resume(resume_path, jd_hard)
            print(f"  Reviewer Pass 3: {c3['score']}/100  (tightened: {applied3})", end="")
            if skipped3:
                print(f"  | remaining: {skipped3}", end="")
            print()
            current_critique = c3

    # SCORING-INTEGRITY INVARIANT: content-QA only. This function NEVER writes
    # to the tracker. The authoritative True ATS score is computed and written
    # solely by the separate ats_score script. (Writing the checklist score here
    # is what inflated every role to 100 in earlier versions.)
    print(f"  Content-QA checklist score: {current_critique['score']}/100  "
          f"(NOT written to tracker)")
    return current_critique["score"]

# =============================================================================
# NOTE: Early-version role content (FairSquare, Proofpoint, REMAX, Aledade)
# was removed from this file on 2026-05-05. Those documents used the v1
# methodology (single-paragraph summary, formulaic cover letters) and were
# pre-differentiation standard. They are preserved in the Submitted/ folder.
# This file now contains only shared functions — no generation logic.
# All role-specific content lives in batch scripts (scripts/build_batch_*.py).
# =============================================================================

