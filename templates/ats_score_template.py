"""
ats_score_TEMPLATE.py — copy to scripts/ats_score_<date>.py for each run.

THIS IS THE AUTHORITATIVE TRACKER SCORE. Run it AFTER build_batch_<date>.py.

It reads each built resume's actual text and scores it against a 25-40 term
jd_full list (weights 3/2/1) built from the real job description. It writes that
True ATS Score into the tracker's "Resume Score" column.

Honesty rule that keeps scores realistic (not auto-100): INCLUDE the genuine
discriminators a sharp recruiter would weigh — required tools you lack, the
company's domain/sector, exact must-have phrases. They register as misses and
pull the score down to where it honestly belongs. A strong fit lands high-80s/
90s; a partial fit lands 70s. A flat 100 means your term list was too soft.
"""
import os, re, importlib.util
try:
    import openpyxl
    from docx import Document  # imported lazily in get_resume_text; checked here for a clear early error
except ImportError as e:
    import sys
    sys.exit(
        f"\n[job-search-tailor] Missing dependency: {e.name}\n"
        "The scorer needs openpyxl (tracker) and python-docx (read resumes).\n"
        "Install them into the SAME Python you run scripts with:\n\n"
        "    python3 -m pip install -r requirements.txt\n\n"
        "Use 'python3 -m pip' (not bare 'pip') so the installer and the\n"
        "interpreter that runs this script are the same one.\n"
    )

# ── resolve the working folder by walking up to the dir that holds .system/ ──
def _find_home(start):
    d = os.path.dirname(os.path.abspath(start))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".system")) and os.path.isdir(os.path.join(d, "profile")):
            return d
        d = os.path.dirname(d)
    raise SystemExit("Could not locate the job-search working folder "
                     "(no .system/ ancestor of this script).")

def _find_tracker(home):
    tdir = os.path.join(home, "tracker")
    files = sorted(f for f in os.listdir(tdir)
                   if f.endswith(".xlsx") and not f.startswith("~$")) if os.path.isdir(tdir) else []
    if not files:
        raise SystemExit("No tracker .xlsx found in <home>/tracker/. Run setup-profile first.")
    return os.path.join(tdir, files[0])

# Resolved at import when run as a real script (this file sits inside a workspace).
# Tolerant of import from outside a workspace (e.g. unit tests) — main() re-resolves
# and raises the clear error then if still unset.
try:
    HOME    = _find_home(__file__)
    OUT     = os.path.join(HOME, "docs", "current") + os.sep   # résumés are read from here
    TRACKER = _find_tracker(HOME)
except SystemExit:
    HOME = OUT = TRACKER = None
SHEET   = "Applications"            # fallback name; the real tab is found by header (see _resolve_sheet)


def get_resume_text(path):
    from docx import Document
    doc = Document(path)
    chunks = [p.text.lower() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text.lower())
    return " ".join(chunks)


def ats_score(resume_text, jd_terms):
    total = sum(w for _, w in jd_terms)
    found_w, found, missed = 0, [], []
    for term, w in jd_terms:
        if re.search(r'\b' + re.escape(term.lower()) + r'\b', resume_text):
            found_w += w; found.append(f"{term}[{w}]")
        else:
            missed.append(f"{term}[{w}]")
    return (round(found_w / total * 100) if total else 0), found, missed


def _resolve_sheet(wb):
    """Locate the Applications tab by its row-3 'Resume Score' header.

    The contract is the header layout, not the tab name — this finds the right
    sheet even when the tab is emoji-prefixed or renamed (the shipped tracker
    uses '📋 Applications'). Falls back to the configured SHEET name, then the
    active sheet; write_tracker() then raises a clear error if the column is
    still missing.
    """
    for ws in wb.worksheets:
        if "Resume Score" in [c.value for c in ws[3]]:
            return ws
    return wb[SHEET] if SHEET in wb.sheetnames else wb.active


def write_tracker(row_num, score):
    """Write score into the 'Resume Score' column of the given row (header on row 3)."""
    wb = openpyxl.load_workbook(TRACKER)
    ws = _resolve_sheet(wb)
    headers = [c.value for c in ws[3]]
    try:
        col = headers.index("Resume Score") + 1
    except ValueError:
        raise SystemExit("Could not find a 'Resume Score' column in row 3 of the tracker.")
    ws.cell(row=row_num, column=col, value=score)
    wb.save(TRACKER)
    print(f"  [tracker → row {row_num}: {score}]")


# ── J1: the mandatory end-of-batch summary table ────────────────────────────
# Emitted here (not free-authored) so every row's Score IS the True ATS Score
# written above — the table and the tracker can't disagree.
def _term_and_weight(formatted):
    """'business intelligence[3]' -> ('business intelligence', 3)."""
    m = re.match(r"^(.*)\[(\d+)\]$", formatted)
    return (m.group(1), int(m.group(2))) if m else (formatted, 0)


def _strongest(formatted_terms):
    """The highest-weight term in a found/missed list, or None."""
    if not formatted_terms:
        return None
    return max((_term_and_weight(t) for t in formatted_terms), key=lambda tw: tw[1])[0]


def summary_row(company, title, base, score, found, missed):
    """A scored table row. 'Why' = strongest matched strength; gap: genuine gap."""
    strength = _strongest(found) or "—"
    gap = _strongest(missed)
    return {"score": score, "company": company, "title": title, "base": base,
            "why": f"{strength}; gap: {gap if gap else 'none'}"}


def deferred_row(company, title, base, reason):
    """A role with no retrievable JD — listed, never silently dropped."""
    return {"score": None, "company": company, "title": title,
            "base": base or "—", "why": f"deferred: {reason}"}


def summary_table(rows):
    """Markdown table sorted by score desc (deferred rows last) + MIN/AVG/MAX."""
    scored = sorted((r for r in rows if r["score"] is not None),
                    key=lambda r: -r["score"])
    deferred = [r for r in rows if r["score"] is None]
    lines = ["| Score | Role | Base | Why |", "|---|---|---|---|"]
    for r in scored + deferred:
        sc = r["score"] if r["score"] is not None else "—"
        lines.append(f"| {sc} | {r['company']} — {r['title']} | {r['base']} | {r['why']} |")
    if scored:
        s = [r["score"] for r in scored]
        stats = f"MIN {min(s)} · AVG {round(sum(s) / len(s))} · MAX {max(s)}"
    else:
        stats = "MIN — · AVG — · MAX —"
    note = ("These are honest True ATS Scores: 70s = partial fit, "
            "high-80s–90s = strong fit.")
    return "\n".join(lines) + f"\n\n**{stats}**  \n_{note}_"


# ════════════════════════════════════════════════════════════════════════════
# ONE BLOCK PER ROLE. Build jd_full from the ACTUAL job description:
#   weight 3 (~5-6): required / must-have terms + the role's defining words
#   weight 2 (~10):  important terms, repeated in the JD, domain-critical
#   weight 1 (~10):  nice-to-haves, supporting vocabulary, sector/domain words
# Include tools/domains you LACK so the score is honest. 25-40 terms total.
# row = the tracker row number you seeded for this job.
# base = "LEADER" or "HANDSON" (matches build_batch) — shown in the summary table.
# A role whose JD couldn't be retrieved: give it {"deferred": "<reason>"} instead
# of resume/jd_full — it's listed in the table with score "—", never dropped.
# ════════════════════════════════════════════════════════════════════════════
ROLES = [
    {
        "row": 4,
        "company": "Example Co",
        "role": "Director of Analytics",
        "base": "LEADER",
        "resume": "001_Your Name - Example Co - Director of Analytics - Resume.docx",
        "jd_full": [
            ("business intelligence", 3), ("analytics", 3), ("data quality", 3),
            ("stakeholder", 3), ("reporting", 3),
            ("kpi", 2), ("executive", 2), ("data governance", 2), ("dashboards", 2),
            ("team leadership", 2), ("data strategy", 2), ("performance", 2),
            ("sql", 1), ("tableau", 1), ("power bi", 1), ("self-service", 1),
            ("transformation", 1), ("mentoring", 1), ("metrics", 1),
            # include genuine gaps so the score stays honest, e.g.:
            ("your_sector_domain", 2), ("a_required_tool_you_lack", 2),
        ],
    },
]

def main():
    global HOME, OUT, TRACKER
    if HOME is None:  # imported tolerantly; resolve now and surface a clear error if not in a workspace
        HOME = _find_home(__file__)
        OUT = os.path.join(HOME, "docs", "current") + os.sep
        TRACKER = _find_tracker(HOME)
    rows = []
    for r in ROLES:
        if r.get("deferred"):
            print(f"\nRow {r.get('row', '—')} | {r['company']} — {r.get('role', '')}: "
                  f"DEFERRED ({r['deferred']}) — not scored")
            rows.append(deferred_row(r["company"], r.get("role", ""),
                                     r.get("base"), r["deferred"]))
            continue
        txt = get_resume_text(OUT + r["resume"])
        score, found, missed = ats_score(txt, r["jd_full"])
        print(f"\nRow {r['row']} | {r['company']} — {r['role']}")
        print(f"  Score:  {score}/100  ({len(found)}/{len(found)+len(missed)} terms)")
        print(f"  Found:  {', '.join(found)}")
        print(f"  Missed: {', '.join(missed)}")
        write_tracker(r["row"], score)
        rows.append(summary_row(r["company"], r["role"], r.get("base", "—"),
                                score, found, missed))

    # The standard batch summary — paste this to the user (J1). Each row's Score
    # is the exact number written to the tracker above.
    print("\n" + "=" * 60)
    print("Batch summary (paste to the user):\n")
    print(summary_table(rows))
    print("=" * 60)


# Run only when executed as a script (e.g. python3 scripts/ats_score_<date>.py).
# Guarding this lets the regression tests import write_tracker()/ats_score()
# without triggering a real scoring run.
if __name__ == "__main__":
    main()
