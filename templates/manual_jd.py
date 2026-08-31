"""manual_jd.py — the manual job-description fallback (Epic W / v0.5.0).

Some postings simply can't be fetched: they 404 after you applied, they render
client-side, they sit behind a login, or the site returns 403. For those, the
user saves the description into

    <home>/docs/manual_job_descriptions/

as a document named with the general strategy

    YYYYMMDD Company - Job Title.docx

and the run reads it instead of asking for a paste. This module does the
deterministic half of that: match a tracker row's Company + Role to a file in
that folder, and read the file's text.

  find_manual_jd()    best-matching document for one tracker row (or None)
  manual_fallback()   the retrieval-path entry point — match + read in one call
  list_manual_jds()   everything parseable in the folder (for reporting)
  stale_manual_jds()  files a PREVIOUS run already consumed
  archive_manual_jds() move used files into archive/ (best-effort, U1)

Matching is by Company + Job Title only; the date prefix is never required and
is used solely to break ties toward the newest file. Nothing here guesses a JD:
a file that matches but can't be text-extracted (e.g. a .pdf) is returned with
`readable=False` so the caller can say so rather than invent a description.
"""
import os
import re

# Where the user drops manual JDs, relative to the job-search working folder.
DIRNAME = os.path.join("docs", "manual_job_descriptions")
# Used JDs are moved here so the drop folder only ever shows what's still live.
ARCHIVE_DIRNAME = "archive"

# Formats we can pull text out of with the plugin's existing dependencies.
READABLE_EXT = (".docx", ".md", ".markdown", ".txt")

# Minimum combined similarity for a file to be accepted as this row's JD.
MIN_SCORE = 0.62
# The company half is a hard gate — a title match at the wrong company is never
# the right document.
MIN_COMPANY_SIM = 0.5

# YYYYMMDD / YYYY-MM-DD / YYYY_MM_DD, only as a *prefix*, and always optional.
_DATE_RE = re.compile(r"^\s*(\d{4})[-_ ]?(\d{2})[-_ ]?(\d{2})\s*[-_ ]\s*")
# Hyphen, en dash, or em dash surrounded by spaces — the company/title divider.
_SEP_RE = re.compile(r"\s+[-–—]\s+")

# Dropped before comparing company names: they're written inconsistently.
_LEGAL = frozenset({
    "inc", "llc", "ltd", "limited", "corp", "corporation", "co", "company",
    "the", "group", "holdings", "holding", "plc", "gmbh", "ag", "nv", "sa",
    "technologies", "technology", "labs",
})
# Dropped before comparing job titles: they carry no discriminating signal.
_TITLE_STOP = frozenset({
    "a", "an", "and", "at", "for", "in", "of", "the", "to", "with",
    "job", "jd", "posting", "description", "role", "position", "opening", "req",
})


# ── folder + filename parsing ───────────────────────────────────────────────
def manual_jd_dir(home):
    """The manual-JD folder for a working folder (not created here)."""
    return os.path.join(home, DIRNAME)


def manual_jd_archive_dir(home):
    """The archive subfolder used JDs are moved into (not created here)."""
    return os.path.join(manual_jd_dir(home), ARCHIVE_DIRNAME)


def parse_manual_filename(name):
    """Split `YYYYMMDD Company - Job Title.docx` into its parts, or None.

    The date prefix is optional and accepts YYYYMMDD / YYYY-MM-DD / YYYY_MM_DD.
    The FIRST ' - ' divides company from title, so a title containing its own
    dash ("Director - Analytics") still parses. A name with no divider returns
    None — it isn't following the strategy, so it is never matched.
    """
    stem, ext = os.path.splitext(os.path.basename(str(name)))
    date = None
    m = _DATE_RE.match(stem)
    if m:
        date = "-".join(m.groups())
        stem = stem[m.end():]
    parts = _SEP_RE.split(stem, maxsplit=1)
    if len(parts) != 2:
        return None
    company, title = parts[0].strip(), parts[1].strip()
    if not company or not title:
        return None
    return {
        "date": date,
        "company": company,
        "title": title,
        "ext": ext.lower(),
        "filename": os.path.basename(str(name)),
    }


# ── similarity ──────────────────────────────────────────────────────────────
def _tokens(text, drop=frozenset()):
    """Lower-cased word tokens with `drop` removed — unless dropping would empty
    the set (a company literally named "The Group" still has to compare)."""
    toks = re.findall(r"[a-z0-9+#]+", str(text or "").lower())
    kept = {t for t in toks if t not in drop}
    return kept or set(toks)


def _company_sim(a, b):
    """1.0 when one company's tokens contain the other's ("Acme" vs "Acme Inc"),
    otherwise Jaccard overlap. Companies get written loosely; be forgiving."""
    if not a or not b:
        return 0.0
    if a <= b or b <= a:
        return 1.0
    return len(a & b) / len(a | b)


def _title_sim(a, b):
    """Overlap over the LARGER title — strict enough that "Director" does not
    fully match "Senior Director of Data", so two roles at the same company
    stay distinguishable."""
    if not a or not b:
        return 0.0
    return len(a & b) / max(len(a), len(b))


def match_score(company, role, parsed):
    """(score, company_sim) for one parsed filename against a tracker row.

    score is the mean of the company and title similarities; company_sim is
    returned separately because it is a hard gate, not just a contributor.
    """
    if not parsed:
        return 0.0, 0.0
    cs = _company_sim(_tokens(company, _LEGAL), _tokens(parsed["company"], _LEGAL))
    ts = _title_sim(_tokens(role, _TITLE_STOP), _tokens(parsed["title"], _TITLE_STOP))
    return (cs + ts) / 2.0, cs


# ── lookup ──────────────────────────────────────────────────────────────────
def list_manual_jds(home):
    """Every parseable document in the manual-JD folder, newest date first.

    Skips Office lock files (`~$…`), dotfiles, and anything that doesn't follow
    the `Company - Job Title` strategy.
    """
    d = manual_jd_dir(home)
    if not os.path.isdir(d):
        return []
    out = []
    for fn in sorted(os.listdir(d)):
        if fn.startswith(("~$", ".")) or not os.path.isfile(os.path.join(d, fn)):
            continue
        parsed = parse_manual_filename(fn)
        if not parsed:
            continue
        parsed["path"] = os.path.join(d, fn)
        parsed["readable"] = parsed["ext"] in READABLE_EXT
        out.append(parsed)
    out.sort(key=lambda p: (p["date"] or "", p["filename"]), reverse=True)
    return out


def find_manual_jd(home, company, role, min_score=MIN_SCORE):
    """The best-matching manual JD for a tracker row, or None.

    Returns the parsed dict plus `path`, `score`, and `readable`. Ties break
    toward the newest dated filename (list_manual_jds is already sorted that
    way, and max() keeps the first best).
    """
    best = None
    for parsed in list_manual_jds(home):
        score, cs = match_score(company, role, parsed)
        if cs < MIN_COMPANY_SIM or score < min_score:
            continue
        if best is None or score > best["score"]:
            parsed["score"] = round(score, 3)
            best = parsed
    return best


# ── reading ─────────────────────────────────────────────────────────────────
def read_manual_jd(path):
    """The plain text of a manual JD document.

    .docx is read via python-docx (paragraphs + table cells, so a JD laid out in
    a table isn't lost); .md/.txt are read as-is. Any other format raises — the
    caller reports it instead of scoring from a guess.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext not in READABLE_EXT:
        raise ValueError(
            f"Cannot extract text from '{os.path.basename(path)}' ({ext}). "
            "Save it as .docx, .md, or .txt.")
    if ext != ".docx":
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(t for t in parts if t and t.strip())


def manual_fallback(home, company, role, min_score=MIN_SCORE):
    """Retrieval-path entry point: find the manual JD for a row and read it.

    Returns None when nothing matches. Otherwise the match dict with `text`
    filled in, or `text=None` plus a `reason` when the matched file exists but
    can't be text-extracted — a matched-but-unreadable file is reported, never
    silently treated as "no JD".
    """
    match = find_manual_jd(home, company, role, min_score)
    if not match:
        return None
    if match["readable"]:
        match["text"] = read_manual_jd(match["path"])
        match["reason"] = None
    else:
        match["text"] = None
        match["reason"] = (f"matched '{match['filename']}' but {match['ext']} "
                           "text can't be extracted — save it as .docx/.md/.txt")
    return match


# ── archiving used JDs (Epic W / U1) ────────────────────────────────────────
def stale_manual_jds(home, active_paths=(), processed=()):
    """Manual JDs a PREVIOUS run already consumed — safe to archive at run start.

    A file is stale when it parses, this batch is not about to use it, and its
    company + title already match a role recorded in the skills-demand index.
    A document the user just dropped in for today has no index row yet, so it is
    never swept up. `processed` is that index: dicts with company/role keys (as
    `skills_demand.load_index()` returns) or (company, role) pairs.
    """
    active = {os.path.abspath(p) for p in active_paths}
    pairs = []
    for item in processed:
        if isinstance(item, dict):
            pairs.append((item.get("company", ""), item.get("role", "")))
        else:
            pairs.append((item[0], item[1]))
    stale = []
    for parsed in list_manual_jds(home):
        if os.path.abspath(parsed["path"]) in active:
            continue
        for company, role in pairs:
            score, cs = match_score(company, role, parsed)
            if cs >= MIN_COMPANY_SIM and score >= MIN_SCORE:
                stale.append(parsed)
                break
    return stale


def archive_manual_jds(home, paths):
    """Move used manual JDs into docs/manual_job_descriptions/archive/.

    Best-effort per U1: a move that is declined or fails never aborts the run.
    Returns (moved, left_behind) where left_behind is [(path, reason), ...] so
    the caller can name what it couldn't move instead of silently piling up.
    A name collision in archive/ is resolved by suffixing, never by overwriting.
    """
    import shutil

    dest_dir = manual_jd_archive_dir(home)
    moved, left = [], []
    for path in paths:
        if not os.path.isfile(path):
            continue
        try:
            os.makedirs(dest_dir, exist_ok=True)
            base, ext = os.path.splitext(os.path.basename(path))
            dest = os.path.join(dest_dir, base + ext)
            n = 2
            while os.path.exists(dest):
                dest = os.path.join(dest_dir, f"{base} ({n}){ext}")
                n += 1
            shutil.move(path, dest)
            moved.append(dest)
        except OSError as e:
            left.append((path, str(e)))
    return moved, left
