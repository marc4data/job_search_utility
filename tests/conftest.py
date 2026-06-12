"""Shared pytest fixtures for the job-search-tailor regression suite.

These tests lock the project's two hard invariants in place:
  • critique_and_refine() is content-QA only and NEVER writes the tracker.
  • The "Resume Score" column is written ONLY by the ats_score script.

The engine reads its PROFILE at import time, so we load a *fresh* copy of
engine/build_docs.py per test, bound to a throwaway dummy profile and a temp
output dir via the JS_PROFILE / JS_OUT_DIR env vars the engine already honors.
"""
import importlib.util
import os

import pytest

REPO_ROOT         = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE_PATH       = os.path.join(REPO_ROOT, "engine", "build_docs.py")
ATS_TEMPLATE_PATH = os.path.join(REPO_ROOT, "templates", "ats_score_template.py")

# A self-contained dummy profile with NO placeholder tokens, so test (b) can
# assert the rendered resume contains only real content.
DUMMY_PROFILE = '''\
PROFILE = {
    "name": "Dana Testwell",
    "contact": "Testville, TS  \\u2022  (555) 000-0000  \\u2022  dana@test.invalid",
    "cl_contact": "Testville, TS  |  (555) 000-0000  |  dana@test.invalid",
    "links": [("LinkedIn", "https://linkedin.invalid/in/dana")],
    "education": ["BS in Testing  \\u2022  Test University  \\u2022  Testville, TS"],
    "certifications": ["Certifications: Test Certified Pro"],
    "additional_experience": ["Junior Tester  \\u2022  Old Test Co, Testville, TS"],
}
'''

# The profile name as it should appear in a built resume (header is upper-cased).
DUMMY_NAME      = "Dana Testwell"
DUMMY_NAME_UPPER = DUMMY_NAME.upper()


def _load_module(name, path):
    """Import a Python file by path as a uniquely-named module."""
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture
def engine(tmp_path, monkeypatch):
    """A fresh build_docs module bound to the dummy profile + a temp out dir."""
    profile_dir = tmp_path / "profile"
    profile_dir.mkdir()
    profile_py = profile_dir / "profile.py"
    profile_py.write_text(DUMMY_PROFILE)

    monkeypatch.setenv("JS_PROFILE", str(profile_py))
    monkeypatch.setenv("JS_OUT_DIR", str(tmp_path))
    # Run with cwd = tmp_path so any *accidental* relative file write during a
    # build/critique would land here, where test (a) can detect it.
    monkeypatch.chdir(tmp_path)
    return _load_module("build_docs_under_test", ENGINE_PATH)


@pytest.fixture
def ats_module():
    """The ats_score template, imported without running its main() loop."""
    return _load_module("ats_score_under_test", ATS_TEMPLATE_PATH)


@pytest.fixture
def docx_text():
    """Return a helper that flattens a .docx into one searchable string."""
    from docx import Document

    def _extract(path):
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    return _extract
